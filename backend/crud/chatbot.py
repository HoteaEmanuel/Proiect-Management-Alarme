from sqlalchemy.orm import Session
from sqlalchemy import select, text
from docx import Document as DocxDocument
from sqlalchemy.exc import IntegrityError, ProgrammingError, SQLAlchemyError
import pandas as pd
import io
import fitz
import json
import requests
import logging

from models import MessageModel, ConversationModel, ConversationFileModel
from schemas import MessageCreate, CloudinaryFileAttachment
from integrations.cloudinary import get_signed_url, delete_files_from_cloudinary
from core import EntityNotFoundError, FileProcessingError, EmptyContentError, InvalidFilenameError, UnsupportedFileFormatError
from core import BaseAppException, DatabaseOperationError, LLMQueryExecutionError

logger = logging.getLogger(__name__)

# Extracts and returns text content from a PDF file byte stream
def parse_pdf(content: bytes) -> str:
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page_index, page in enumerate(doc):
            text=page.get_text("text") or ""
            if text.strip():
                pages.append(
                    f"\n\n[Page {page_index+1}]\n{text.strip()}"
                )
        extracted_text="\n".join(pages).strip()
        
    except Exception as e:
        logger.error(f"PDF internal extraction error: {str(e)}", exc_info=True)
        raise FileProcessingError("The PDF file is corrupted or the format is not supported.")
    
    if not extracted_text:
            raise EmptyContentError("The PDF document was processed successfully, but contains no text")
    
    return extracted_text
    
# Parses DOCX files
def parse_docx(content:bytes) -> str:
    try:
        doc=DocxDocument(io.BytesIO(content))
        paragraphs=[p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells=[cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        extracted = "\n".join(paragraphs).strip()
        
    except Exception as e:
        logger.error(f"Docx internal extraction error: {str(e)}", exc_info=True)
        raise FileProcessingError("The docx file is corrupted or the format is not suported.")
    
    if not extracted:
            raise EmptyContentError("The docx document was processed successfully, but contains no text")
    
    return extracted

# Parses a file with a specific extension
def parse_file(filename: str, content: bytes) -> str:
    if "." not in filename:
        raise InvalidFilenameError()
    
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext=="pdf":
        return parse_pdf(content)
    elif ext=="docx":
        return parse_docx(content)
    elif ext in ("csv", "xlsx", "xls"):
        try:
            if ext == "csv":
                df=pd.read_csv(io.BytesIO(content), on_bad_lines="skip")
                return df.to_json(orient="records", force_ascii=False)
            else:
                df=pd.read_excel(io.BytesIO(content))
                return df.to_markdown(index=False)
        except Exception as e:
            logger.error(f"Pandas parsing error for file {filename}: {str(e)}", exc_info=True)
            raise FileProcessingError(f"Could not read the {ext.upper()} file. The content might be corrupted.")
    else:
        raise UnsupportedFileFormatError(ext)

# Retrieves a list of files associated with a specific message ID
def get_message_files(db: Session, message_id: int) -> list[dict]:
    rows = db.execute(
        select(ConversationFileModel)
        .where(
            ConversationFileModel.message_id == message_id,
            ConversationFileModel.is_deleted == False
        )
    ).scalars().all()

    return [
        {
            "filename": f.filename,
            "url": f.file_url,
            "public_id": f.public_id,
            "resource_type": f.resource_type,
            "file_format": f.file_format,
            "file_size": f.file_size,
        }
        for f in rows
    ]

# Parses the raw JSON content of an assistant's message and extracts the combined text from its blocks
def _parse_assistant_content(raw_content: str) -> str:
    try:
        blocks = json.loads(raw_content)
        return " ".join(
            block["content"] for block in blocks
            if block.get("type") == "text"
        )
    except (json.JSONDecodeError, TypeError, KeyError):
        return raw_content

# Retrieves, downloads, and parses files attached to a user message, appending their content to the base text
def _parse_user_files(db: Session, message_id: str, base_content: str) -> str:
    files = get_message_files(db, message_id)
    if not files:
        return base_content
    
    parsed_files = []
    for file in files:
        filename = file["filename"]
        try:
            signed_url = get_signed_url(file["public_id"], file["resource_type"])
            logger.debug(f"[SIGNED URL] {signed_url}")

            response = requests.get(signed_url, allow_redirects=True, timeout=15)
            response.raise_for_status() #Raise error if HTTP status is not 2xx (expired URL)

            parsed_text = parse_file(filename, response.content)
            parsed_files.append(f"[{filename}]:\n{parsed_text}")
        except BaseAppException as app_error:
            logger.warning(f"[CONTEXT FILES] Error parsing '{filename}': {app_error.message}")
            continue
        except Exception as e:
            logger.error(f"[CONTEXT FILES] Unexpected error when downlading/parsing '{filename}': {str(e)}")
            continue

    if parsed_files:
        return base_content + "\n\nAttached files:\n\n" + "\n\n".join(parsed_files)
    
    return base_content

# Retrieves and parses a limited history of messages for a specific conversation, returning them in chronological order
def get_conversation_history(db: Session, user_id: str, conversation_id: str, limit: int = 10) -> list[dict]:
    stmt = (
        select(MessageModel)
        .where(
            MessageModel.user_id == user_id,
            MessageModel.conversation_id == conversation_id
        )
        .order_by(MessageModel.created_at.desc())
        .limit(limit)
    )
    
    rows = db.execute(stmt).scalars().all()
    messages = list(reversed(rows))
    
    result = []
    for msg in messages:
        if msg.role == "assistant":
            content = _parse_assistant_content(msg.content)
        else:
            content = _parse_user_files(db, msg.id, msg.content)
        
        result.append({
            "role": msg.role,
            "content": content
        })

    return result

# Retrieves a list of all uploaded files within a specific conversation
def get_file_history(db: Session, user_id: str, conversation_id: str) -> list[CloudinaryFileAttachment]:
    files = (
        db.query(ConversationFileModel, MessageModel.role)
        .join(MessageModel, MessageModel.id == ConversationFileModel.message_id)
        .filter(
            MessageModel.conversation_id == conversation_id,
            MessageModel.user_id == user_id,
            ConversationFileModel.is_deleted == False
        )
        .order_by(ConversationFileModel.created_at.asc())
        .all()
    )

    return [
        CloudinaryFileAttachment(
            filename=f.filename,
            url=f.file_url,
            public_id=f.public_id or "",
            resource_type=f.resource_type or "",
            file_format=f.file_format or "",
            file_size=f.file_size or 0,
            role=role
        )
        for f, role in files
    ]

# Fetches a specific conversation record for a given user
def get_conversation_data(db: Session, user_id: str, conversation_id: str) -> ConversationModel:
    conversation = db.execute(
        select(ConversationModel)
        .where(
            ConversationModel.conversation_id == conversation_id,
            ConversationModel.user_id == user_id
        )
    ).scalar()

    if conversation is None:
        raise EntityNotFoundError(f"Could not find conversation!")

    return conversation
        

def _parse_assistant_blocks(raw_content: str):
    try:
        return json.loads(raw_content)
    except (json.JSONDecodeError, TypeError, ValueError):
        return raw_content

# Retrives the history of a conversations
def get_full_conversation(db: Session, user_id: str, conversation_id: str):
    conversation = db.execute(
        select(ConversationModel)
        .where(
            ConversationModel.conversation_id == conversation_id,
            ConversationModel.user_id == user_id
        )
    ).scalar()

    if conversation is None:
        raise EntityNotFoundError("Conversation", conversation_id)

    stmt = (
        select(MessageModel)
        .where(
            MessageModel.user_id == user_id,
            MessageModel.conversation_id == conversation_id
        )
        .order_by(MessageModel.created_at.asc())
    )
    messages = db.execute(stmt).scalars().all()
    
    result = []
    for msg in messages:
        if msg.role == "assistant":
            entry = {
                "role": msg.role,
                "blocks": _parse_assistant_blocks(msg.content),
                "is_stopped": msg.is_stopped,
                "smart_replies": msg.smart_replies
            }
        else:
            entry = {
                "role": msg.role,
                "content": msg.content
            }
            
        files = get_message_files(db, msg.id)
        if files:
            entry["files"] = files

        result.append(entry)

    return result

# Retrieves a list of all conversations belonging to a specific user
def get_user_conversations(db: Session, user_id: str) -> list[ConversationModel]:
    stmt = (
        select(ConversationModel)
        .where(
            ConversationModel.user_id == user_id
        )
        .order_by(ConversationModel.created_at.desc())
    )

    rows = db.execute(stmt).scalars().all()
    
    return rows

# Saves a new message to the database
def save_message_to_db(db: Session, message_data: MessageCreate) -> MessageModel:
    
    message = MessageModel(
        conversation_id=message_data.conversation_id,
        user_id=message_data.user_id,
        role=message_data.role,
        content=message_data.content,
        smart_replies=message_data.smart_replies,
        has_sql_query=message_data.has_sql_query,
        sql_query=message_data.sql_query,
        is_stopped=message_data.is_stopped
    )
    
    try:
        db.add(message)
        db.commit()
        db.refresh(message)
        return message
    
    except IntegrityError as e:
        db.rollback()
        logger.error(f"Integrity error saving message {message_data}: {str(e)}")
        raise DatabaseOperationError("Could not save the message due to a data conflict.")
    

# Updates a user's message record to link it with the corresponding bot response ID
def set_response_id(db: Session, user_message_id: int, bot_response_id: int) -> None:
    updated_rows = db.query(MessageModel)\
        .filter(MessageModel.id == user_message_id)\
        .update({"response_id": bot_response_id})

    if updated_rows == 0:
        raise EntityNotFoundError("user_message", user_message_id)

    try:
        db.commit()
    except IntegrityError as e:
        db.rollback()
        logger.error(f"Integrity error updating user message {user_message_id}: {str(e)}")
        raise DatabaseOperationError("Could not update the user message response due to a data conflict.")

# Retrieves the title of a conversation based on its ID
def get_conversation_title(db: Session, conversation_id: str) -> str:
    
    result = db.execute(
        text("SELECT CONVERSATION_TITLE FROM CONVERSATIONS WHERE CONVERSATION_ID = :conversation_id"),
        {"conversation_id": conversation_id}
    ).scalar()

    if result is None:
        raise EntityNotFoundError("Conversation", conversation_id)
    
    return result

# Updates the title of a specific conversation in the database
def set_conversation_title(db: Session, conversation_id: str, conversation_title: str) -> None:
    updated_rows = db.query(ConversationModel)\
        .filter(ConversationModel.conversation_id == conversation_id)\
        .update({"conversation_title": conversation_title})

    if updated_rows == 0:
        raise EntityNotFoundError("Conversation", conversation_id)

    try:
        db.commit()

    except IntegrityError as e:
        db.rollback()
        logger.error(f"Integrity error updating conversation {conversation_id}: {str(e)}")
        raise DatabaseOperationError("Could not update the conversation title due to a data conflict.")
# Creates a new empty conversation for the specified user
def create_new_conversation(db: Session, user_id: str) -> ConversationModel:
    
    conversation = ConversationModel(
        user_id=user_id,
    )
    
    try:
        db.add(conversation)
        db.commit()
        db.refresh(conversation)

    except IntegrityError as e:
        db.rollback()
        logger.error(f"Could not save the new conversation to the database: {str(e)}")
        raise DatabaseOperationError("Could not save the new conversation.")
    
    return conversation

# Executes a raw SQL query generated by the LLM
def run_llm_query(db: Session, query: str) -> list[dict]:
    try:
        result = db.execute(text(query))
        if result.returns_rows:
            return [dict(row) for row in result.mappings().all()]
        
        return []
    except ProgrammingError as e:
        db.rollback()
        logger.error(f"The LLM generated invalid SQL syntax. Query: {query} | Error: {str(e)}")
        raise LLMQueryExecutionError("The query generated by the LLM has invalid syntax.")
    
    except SQLAlchemyError as e:
        db.rollback()
        logger.error(f"Error caused by the execution of LLM generated query: {str(e)}")
        raise LLMQueryExecutionError("The query generated by the LLM has caused an error.")
    
    
# Deletes a conversation and all its associated messages    
def delete_conversation(db: Session, user_id: str, conversation_id: str) -> None:
    try:
        conversation=db.execute(
            select(ConversationModel).
            where(
                ConversationModel.conversation_id == conversation_id,
                ConversationModel.user_id == user_id)
        ).scalar()

        if conversation is None:
            raise EntityNotFoundError("Conversation", conversation_id)
        
        files = get_file_history(db, user_id, conversation_id)

        delete_files_from_cloudinary(files)
        
        db.delete(conversation)
        db.commit()
        
    except EntityNotFoundError:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to delete conversation {conversation_id}: {str(e)}")
        raise DatabaseOperationError("An error occurred while deleting the conversation from the database.")

# Modifies the title of an existing conversation
def update_conversation_title(db: Session, user_id: str, conversation_id: str, new_title: str) -> None:
    conversation = db.execute(
        select(ConversationModel)
        .where(
            ConversationModel.conversation_id == conversation_id,
            ConversationModel.user_id == user_id
        )
    ).scalar()

    if conversation is None:
        raise EntityNotFoundError("Conversation", conversation_id)
    
    conversation.conversation_title=new_title
    
    try:
        db.commit()
        db.refresh(conversation)

        return conversation
    
    except IntegrityError as e:
        db.rollback()
        logger.error(f"Integrity error while saving the conversation title: {str(e)}")
        raise DatabaseOperationError("Could not save conversation title due to a data conflict")
    
# Saves metadata for multiple uploaded files to the database
def save_conversation_files(db: Session, message_id: int, files: list[CloudinaryFileAttachment]) -> list[CloudinaryFileAttachment]:
    db_files = [
        ConversationFileModel(
            message_id=message_id,
            filename=file.filename,
            file_url=file.url,
            public_id=file.public_id,
            resource_type=file.resource_type,
            file_format=file.file_format,
            file_size=file.file_size,
        )
        for file in files
    ]
    try:
        db.add_all(db_files)
        db.commit()
        for f in db_files:
            db.refresh(f)
    
        return db_files
    
    except IntegrityError as e:
        db.rollback()
        logger.error(f"Integrity error while saving the message 'persist' files: {str(e)}")
        raise DatabaseOperationError("Could not save the attached files due to a data conflict.")

